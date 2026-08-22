"""Markdown → Word(.docx)。

优先用 pandoc（项目 vendor/ 或 PATH 中的单文件二进制），
找不到时回退到 python-docx 的简易转换。
"""
from __future__ import annotations

import logging
import os
import re
import shutil
import subprocess
from pathlib import Path
from typing import Optional

log = logging.getLogger(__name__)

_PROJECT_ROOT = Path(__file__).resolve().parent.parent


def find_pandoc() -> Optional[str]:
    # 1) 项目 vendor 目录
    vendor_bins = sorted((_PROJECT_ROOT / "vendor").glob("pandoc-*/bin/pandoc"))
    if vendor_bins:
        return str(vendor_bins[-1])
    # 2) PATH
    return shutil.which("pandoc")


def md_to_docx(md_path: str, docx_path: str) -> bool:
    """用 pandoc 转换；成功返回 True。"""
    pandoc = find_pandoc()
    if not pandoc:
        return False
    cmd = [
        pandoc,
        md_path,
        "-o", docx_path,
        "--from", "markdown+yaml_metadata_block+footnotes+smart",
        "--toc",
    ]
    r = subprocess.run(cmd, capture_output=True, text=True)
    if r.returncode != 0:
        log.warning("pandoc 转换失败：%s", r.stderr[:500])
        return False
    return True


def docx_fallback(md_path: str, docx_path: str) -> None:
    """没有 pandoc 时的简易转换：识别 #/##/### 与引用块。"""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    for raw in lines:
        line = raw.rstrip("\n").rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:])
            p.style = doc.styles["Quote"] if "Quote" in doc.styles else doc.styles["Normal"]
        elif line.startswith("<!--"):
            continue  # 页码注释，Word 里不显示
        else:
            doc.add_paragraph(line)

    # 默认正文样式：中文可读字号
    for s in ("Normal", "Body Text"):
        if s in doc.styles:
            doc.styles[s].font.size = Pt(12)
    doc.save(docx_path)


def convert(md_path: str, docx_path: str) -> str:
    """入口：返回使用的转换器名称（'pandoc' / 'python-docx'）。"""
    if md_to_docx(md_path, docx_path):
        return "pandoc"
    docx_fallback(md_path, docx_path)
    return "python-docx"
